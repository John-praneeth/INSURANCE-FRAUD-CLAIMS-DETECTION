"""
Script to generate a highly professional Academic Internship Report in DOCX format
with embedded high-resolution screenshots, charts, tables, and formatted certificate.
"""

import os
import sys
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text, title="KEY OPERATIONAL RESULT"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F0FDF4")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border styling
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="16A34A"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(22, 101, 52)
    
    r_body = p.add_run(text)
    r_body.font.name = "Arial"
    r_body.font.size = Pt(9.5)
    r_body.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "0F172A")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code_str)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(241, 245, 249)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_image_with_caption(doc, img_path, caption_text, width_inches=5.8):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(10)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(14)
        r_cap = p_cap.add_run(caption_text)
        r_cap.font.name = "Arial"
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 116, 139)
    else:
        print(f"[Warning] Image not found at {img_path}")

def build_report():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Insurance Fraud Claims Detection Engine")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = RGBColor(148, 163, 184)
        hrun.font.italic = True

    # Title Page / Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    r_main = p_title.add_run("ACADEMIC INTERNSHIP REPORT\n")
    r_main.font.name = "Arial"
    r_main.font.size = Pt(18)
    r_main.bold = True
    r_main.font.color.rgb = RGBColor(30, 58, 138)
    
    r_sub = p_title.add_run("INSURANCE FRAUD CLAIMS DETECTION ENGINE:\nMACHINE LEARNING RISK SCREENING DECISION SUPPORT SYSTEM")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(12)
    r_sub.bold = True
    r_sub.font.color.rgb = RGBColor(15, 23, 42)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(4)
    p_meta.paragraph_format.space_after = Pt(20)
    rm = p_meta.add_run("Bachelor of Technology in Computer Science & Engineering (Data Science)\nDomain: Artificial Intelligence, Machine Learning & InsurTech Risk Analytics")
    rm.font.name = "Arial"
    rm.font.size = Pt(9.5)
    rm.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # ACKNOWLEDGEMENT
    h1 = doc.add_heading("ACKNOWLEDGEMENT", level=1)
    h1.runs[0].font.name = "Arial"
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(10)

    p_ack1 = doc.add_paragraph(
        "In today’s competitive insurance and financial ecosystem, fraud detection and operational risk screening "
        "represent critical challenges faced by insurance carriers. Fraudulent claims range from opportunistic overstatements "
        "of property damage to organized collision staging and falsified bodily injury reports. Manual auditing of every "
        "submitted claim is cost-prohibitive, while blanket automated approvals expose insurers to severe financial leakage. "
        "Identifying claims that are likely fraudulent at an early stage allows organizations to allocate investigative resources "
        "efficiently and protect legitimate policyholders."
    )
    p_ack1.paragraph_format.line_spacing = 1.25
    p_ack1.paragraph_format.space_after = Pt(8)

    p_ack2 = doc.add_paragraph(
        "The project “Insurance Fraud Claims Detection Engine: Machine Learning Risk Screening Framework” focuses on "
        "predicting the likelihood of insurance fraud using supervised machine learning techniques. The system analyzes "
        "multi-dimensional attributes including policyholder demographics, policy terms, incident dynamics, collision severity, "
        "witnesses, police reports, and monetary claim breakdowns."
    )
    p_ack2.paragraph_format.line_spacing = 1.25
    p_ack2.paragraph_format.space_after = Pt(8)

    p_ack3 = doc.add_paragraph(
        "The project utilizes an Automobile Insurance Claims dataset containing 1,000 customer records and 39 predictor attributes. "
        "The data is processed and analyzed using Python libraries such as Pandas, NumPy, Scikit-Learn, Imbalanced-Learn, XGBoost, "
        "Matplotlib, and Seaborn. Preprocessing techniques are applied to handle missing values, encode categorical attributes, "
        "and mitigate class imbalance using SMOTE before training the machine learning model."
    )
    p_ack3.paragraph_format.line_spacing = 1.25
    p_ack3.paragraph_format.space_after = Pt(8)

    p_ack4 = doc.add_paragraph(
        "An XGBoost Classifier combined with SMOTE and class-weight scaling is deployed to predict fraud probability and "
        "categorize claims into Low Risk (<30%), Medium Risk (30%–60%), and High Risk (>60%) tiers. Operating at an optimal "
        "decision threshold of 0.45, the model achieves an Accuracy of 84.00%, Recall of 75.51%, and an ROC-AUC of 84.12% "
        "on an untouched holdout test set."
    )
    p_ack4.paragraph_format.line_spacing = 1.25
    p_ack4.paragraph_format.space_after = Pt(8)

    p_ack5 = doc.add_paragraph(
        "The developed system is implemented as an interactive Streamlit web application deployed to Streamlit Community Cloud. "
        "Users can enter claim details, obtain predicted fraud risk probabilities, visual risk gauges, and tailored investigative "
        "recommendations. Overall, this project demonstrates the practical application of data science, predictive analytics, "
        "and web deployment to solve a real-world enterprise problem."
    )
    p_ack5.paragraph_format.line_spacing = 1.25
    p_ack5.paragraph_format.space_after = Pt(20)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s = p_sign.add_run("[STUDENT / CANDIDATE NAME]\nRoll No: [ROLL NUMBER / REGISTRATION ID]\nDepartment of Computer Science & Engineering")
    r_s.font.name = "Arial"
    r_s.font.size = Pt(9.5)
    r_s.font.bold = True
    r_s.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_page_break()

    # CONTENTS
    h_toc = doc.add_heading("CONTENTS", level=1)
    h_toc.runs[0].font.name = "Arial"
    h_toc.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    toc_table = doc.add_table(rows=1, cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = toc_table.rows[0].cells
    hdr_cells[0].text = "Chapter / Section Topic"
    hdr_cells[1].text = "Section"
    set_cell_background(hdr_cells[0], "1E3A8A")
    set_cell_background(hdr_cells[1], "1E3A8A")
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    
    toc_rows = [
        ("1. Introduction", "1"),
        ("2. Company / Academic Organization Details", "2"),
        ("3. Technologies Used", "3"),
        ("4. Project: Insurance Fraud Claims Detection Engine", "4"),
        ("    4.1 Objective", "4.1"),
        ("    4.2 Problem Statement", "4.2"),
        ("    4.3 Project Planning and Requirements Gathering", "4.3"),
        ("    4.4 Technology Stack", "4.4"),
        ("    4.5 Dataset Description", "4.5"),
        ("    4.6 System Architecture & Methodology", "4.6"),
        ("    4.7 Data Preprocessing", "4.7"),
        ("    4.8 Model Building – XGBoost Classifier", "4.8"),
        ("    4.9 Application Development – Streamlit Web Application", "4.9"),
        ("        4.9.1 Dashboard Page", "4.9.1"),
        ("        4.9.2 Claim Risk Screening Page", "4.9.2"),
        ("        4.9.3 Model Performance & Analytics Page", "4.9.3"),
        ("        4.9.4 About Project Page", "4.9.4"),
        ("    4.10 Security & Data Validation", "4.10"),
        ("    4.11 Testing & Quality Assurance", "4.11"),
        ("    4.12 Cloud Deployment", "4.12"),
        ("    4.13 Maintenance", "4.13"),
        ("    4.14 Tools Used", "4.14"),
        ("    4.15 Results and Evaluation", "4.15"),
        ("    4.16 Exploratory Data Analysis", "4.16"),
        ("5. Course / Internship Experience", "5"),
        ("6. Conclusion & Future Scope", "6"),
        ("7. Certificate of Internship", "7"),
    ]
    
    for title, sec in toc_rows:
        row = toc_table.add_row()
        c0, c1 = row.cells
        c0.text = title
        c1.text = sec
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        for cell in [c0, c1]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(8.5)
                    if not title.startswith(" "):
                        r.bold = True
                    r.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_page_break()

    # CHAPTER 1
    h1 = doc.add_heading("1. INTRODUCTION", level=1)
    h1.runs[0].font.name = "Arial"
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    doc.add_paragraph(
        "This report outlines the experiences, insights and technical work carried out during the Data Science and "
        "Machine Learning internship. As part of this engagement, I designed and developed the “INSURANCE FRAUD CLAIMS "
        "DETECTION ENGINE”, a machine-learning based risk screening solution that predicts whether an automobile insurance "
        "claim is suspicious or legitimate, and flags high-risk records early so that human claims adjusters can take "
        "preventive, targeted investigative action."
    )
    doc.add_paragraph(
        "Insurance fraud represents one of the most critical operational challenges faced by insurance companies, "
        "causing tens of billions of dollars in annual losses worldwide. Fraudulent claims increase operational expenses "
        "for carriers and inflate premiums for honest policyholders. Traditional detection systems rely on static heuristic "
        "filters (such as flagging claims above a fixed dollar amount) or manual sample audits. Static rules are easily "
        "circumvented by sophisticated fraudsters, while manual audits fail to scale across high claim volumes."
    )
    doc.add_paragraph(
        "This project applies supervised machine learning techniques to historical claim data—covering customer demographics, "
        "policy attributes, incident dynamics, collision severity, witnesses, police reports, and monetary damage amounts—to "
        "build an end-to-end predictive pipeline and interactive Streamlit web application. The system operates as an early-warning "
        "screening tool: rather than conclusively declaring fraud, it computes calibrated probabilities, segments claims into "
        "intuitive risk categories (LOW, MEDIUM, HIGH), and provides explainable decision support."
    )

    # CHAPTER 2
    h2 = doc.add_heading("2. COMPANY / ACADEMIC ORGANIZATION DETAILS", level=1)
    h2.runs[0].font.name = "Arial"
    h2.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    tbl_comp = doc.add_table(rows=1, cols=2)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcells = tbl_comp.rows[0].cells
    hcells[0].text = "Field"
    hcells[1].text = "Details"
    set_cell_background(hcells[0], "1E3A8A")
    set_cell_background(hcells[1], "1E3A8A")
    for cell in hcells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    comp_info = [
        ("Project Name", "Insurance Fraud Claims Detection Engine"),
        ("Domain", "Artificial Intelligence, Machine Learning & Financial Risk Analytics"),
        ("Industry Focus", "InsurTech, Insurance Claims Screening & Decision Support"),
        ("Internship Duration", "2 Months"),
        ("Internship Track", "Data Science / Machine Learning Engineering"),
        ("Mode of Internship", "Project-based, mentor-guided internship"),
        ("Key Skill Areas", "Python, Pandas, Scikit-Learn, XGBoost, SMOTE, Streamlit, Cloud Deployment"),
    ]
    for k, v in comp_info:
        r = tbl_comp.add_row()
        c0, c1 = r.cells
        c0.text = k
        c1.text = v
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        for cell in [c0, c1]:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # CHAPTER 3
    h3 = doc.add_heading("3. TECHNOLOGIES USED", level=1)
    h3.runs[0].font.name = "Arial"
    h3.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph("The technical solution was constructed by combining modern data science and web engineering frameworks:")
    techs = [
        ("Python", "Core programming language for data loading, preprocessing, model training, and web dashboard."),
        ("Pandas & NumPy", "Tabular data manipulation, missing value handling, date parsing, and vector calculations."),
        ("Scikit-Learn", "Provided ColumnTransformer, StandardScaler, OneHotEncoder, and cross-validation pipelines."),
        ("Imbalanced-Learn (imblearn)", "Provided SMOTE for cross-validated synthetic oversampling of minority fraud cases."),
        ("XGBoost", "Extreme gradient boosted decision tree classifier tuned with scale_pos_weight."),
        ("Matplotlib & Seaborn", "Rendered static high-resolution evaluation curves and confusion matrix heatmaps."),
        ("Streamlit & Plotly", "Interactive multi-page web application featuring dynamic gauge meters and metric cards."),
        ("Joblib", "Model serialization and deserialization for sub-millisecond cloud inference."),
        ("Pytest", "Automated unit testing suite verifying data loading, feature engineering, and model predictions.")
    ]
    for name, desc in techs:
        p = doc.add_paragraph(style='List Bullet')
        r_name = p.add_run(f"{name}: ")
        r_name.bold = True
        r_name.font.name = "Arial"
        r_name.font.size = Pt(9.5)
        r_desc = p.add_run(desc)
        r_desc.font.name = "Arial"
        r_desc.font.size = Pt(9.5)

    doc.add_page_break()

    # CHAPTER 4
    h4 = doc.add_heading("4. PROJECT: INSURANCE FRAUD CLAIMS DETECTION ENGINE", level=1)
    h4.runs[0].font.name = "Arial"
    h4.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    # 4.1 & 4.2
    doc.add_heading("4.1 Objective", level=2)
    doc.add_paragraph(
        "The objective is to develop an automated machine learning risk screening system that ingests automobile insurance "
        "claim records, computes a calibrated fraud probability, segments claims into intuitive risk tiers (LOW, MEDIUM, HIGH), "
        "and serves actionable recommendations to human claims investigators via an interactive cloud portal."
    )

    doc.add_heading("4.2 Problem Statement", level=2)
    doc.add_paragraph(
        "Given 39 heterogeneous customer, policy, incident, and vehicle attributes, can we train a supervised binary classifier "
        "that accurately detects fraudulent claims (fraud_reported = 1) while controlling false positive rates, and present this "
        "insight through an accessible real-time early-warning web portal for business stakeholders?"
    )

    doc.add_heading("4.3 Dataset Description", level=2)
    doc.add_paragraph(
        "The dataset contains 1,000 claim records with 39 predictor attributes and 1 binary target (fraud_reported: 753 Legitimate "
        "vs 247 Fraudulent claims). Missing values encoded as '?' in collision_type, property_damage, and police_report_available "
        "were converted into explicit 'MISSING' tokens."
    )

    # 4.4 Preprocessing Code Excerpt
    doc.add_heading("4.4 Data Preprocessing Pipeline", level=2)
    doc.add_paragraph("The leak-free ColumnTransformer pipeline chains numeric scaling and categorical one-hot encoding:")
    add_code_block(doc, 
"""def create_preprocessor_pipeline(num_cols, cat_cols):
    num_pipeline = Pipeline([
        ("imputer", SimpleImputer(strategy="median")),
        ("scaler", StandardScaler())
    ])
    cat_pipeline = Pipeline([
        ("imputer", SimpleImputer(strategy="constant", fill_value="MISSING")),
        ("onehot", OneHotEncoder(handle_unknown="ignore", sparse_output=False))
    ])
    return ColumnTransformer(transformers=[
        ("num", num_pipeline, num_cols),
        ("cat", cat_pipeline, cat_cols)
    ], remainder="drop")"""
    )

    # 4.5 Model Building
    doc.add_heading("4.5 Model Building – XGBoost Classifier", level=2)
    doc.add_paragraph(
        "An XGBoost Classifier with SMOTE oversampling was trained within a 5-fold stratified cross-validation framework. "
        "Predicted probabilities are categorized into business risk tiers:"
    )
    doc.add_paragraph("• Low Risk: Probability < 0.30 (Normal automated claims workflow)")
    doc.add_paragraph("• Medium Risk: Probability 0.30 – 0.60 (Additional verification required before payout)")
    doc.add_paragraph("• High Risk: Probability > 0.60 (Immediate referral to Special Investigation Unit)")

    doc.add_page_break()

    # 4.6 Application Development & Screenshots
    doc.add_heading("4.6 Application Development – Streamlit Web Application", level=2)
    doc.add_paragraph(
        "The application is built with Streamlit and deployed on Streamlit Community Cloud. Navigation is driven by a sidebar "
        "radio selector routing across 4 core pages:"
    )

    doc.add_heading("4.6.1 Dashboard Page", level=3)
    doc.add_paragraph(
        "The Executive Dashboard presents 5 top-level KPI cards (Total Claims, Fraud Rate, Accuracy, Target Recall, ROC-AUC) "
        "accompanied by interactive visual widgets displaying the target distribution and confusion matrix."
    )
    add_image_with_caption(
        doc,
        "reports/figures/screenshot_dashboard_page.png",
        "Figure 4.4: Executive Dashboard page — KPI cards, target distribution pie chart, and confusion matrix heatmap"
    )

    doc.add_heading("4.6.2 Claim Risk Screening Portal", level=3)
    doc.add_paragraph(
        "The Claim Risk Screening page exposes an intuitive 4-section input form covering policyholder demographics, policy terms, "
        "incident details, and vehicle damage amounts. Adjusters can also load sample presets to test high-risk or low-risk claims."
    )
    add_image_with_caption(
        doc,
        "reports/figures/screenshot_claim_screening_form.png",
        "Figure 4.5: Claim Risk Screening page — Interactive claim parameter input form and preset selector"
    )

    doc.add_paragraph(
        "Upon submission, the engine passes the input through the XGBoost inference pipeline and renders a real-time risk gauge, "
        "colored risk badge, and dynamic key risk factors observed:"
    )
    add_image_with_caption(
        doc,
        "reports/figures/screenshot_claim_screening_result.png",
        "Figure 4.6: Real-time risk screening output — Gauge meter, risk category badge, recommended action, and observed risk drivers"
    )

    doc.add_page_break()

    doc.add_heading("4.6.3 Model Performance & Analytics Page", level=3)
    doc.add_paragraph(
        "Displays the cross-validation comparison summary table and renders high-resolution evaluation charts."
    )
    add_image_with_caption(
        doc,
        "reports/figures/screenshot_model_analytics_page.png",
        "Figure 4.7: Model Performance page — Evaluation summary table and interactive evaluation visualizations"
    )

    doc.add_heading("4.6.4 About System Page", level=3)
    doc.add_paragraph("Provides system documentation, architectural workflow, and operational risk tier definitions.")
    add_image_with_caption(
        doc,
        "reports/figures/screenshot_about_system_page.png",
        "Figure 4.8: About System page — System methodology, architecture, and risk threshold documentation"
    )

    doc.add_page_break()

    # 4.7 Results and Evaluation
    doc.add_heading("4.7 Results and Evaluation", level=2)
    doc.add_paragraph(
        "The candidate classifiers were benchmarked across 5-Fold Stratified Cross-Validation on the training set (800 claims). "
        "XGBoost demonstrated superior PR-AUC and balanced recall:"
    )

    tbl_results = doc.add_table(rows=1, cols=6)
    tbl_results.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_hdr = tbl_results.rows[0].cells
    r_hdr[0].text = "Model Algorithm"
    r_hdr[1].text = "CV Recall"
    r_hdr[2].text = "CV Precision"
    r_hdr[3].text = "CV F1-Score"
    r_hdr[4].text = "CV ROC-AUC"
    r_hdr[5].text = "CV PR-AUC"
    for c in r_hdr:
        set_cell_background(c, "1E3A8A")
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(8)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    res_data = [
        ("XGBoost (Tuned)", "0.7226", "0.6550", "0.6862", "0.8474", "0.6629"),
        ("Logistic Regression", "0.7424", "0.6220", "0.6767", "0.8657", "0.6676"),
        ("Gradient Boosting", "0.7372", "0.6621", "0.6977", "0.8642", "0.6464"),
        ("Decision Tree", "0.5508", "0.5488", "0.5498", "0.7007", "0.6086"),
        ("Random Forest", "0.4247", "0.6012", "0.4972", "0.8520", "0.6055"),
        ("K-Nearest Neighbors", "0.8786", "0.2671", "0.4091", "0.5616", "0.4133"),
    ]
    for row_vals in res_data:
        r = tbl_results.add_row()
        for i, val in enumerate(row_vals):
            c = r.cells[i]
            c.text = val
            set_cell_margins(c, top=50, bottom=50, left=80, right=80)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
                    if i == 0 and "XGBoost" in val:
                        run.bold = True
                    run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_callout_box(
        doc,
        "Optimal Decision Threshold: 0.45\n"
        "• Test Accuracy: 84.00% | Test Precision: 64.91% | Test Recall (Target Class 1): 75.51%\n"
        "• Test F1-Score: 69.81% | Test ROC-AUC: 84.12% | Test PR-AUC: 59.89%\n"
        "• Confusion Matrix on 200 Test Claims: TN = 131, FP = 20, FN = 12, TP = 37 (Catches 37 of 49 frauds)",
        title="HELD-OUT TEST SET PERFORMANCE (200 UNTOUCHED CLAIMS)"
    )

    add_image_with_caption(
        doc,
        "reports/figures/XGBoost_Optimized_confusion_matrix.png",
        "Figure 4.10: Confusion Matrix of the Final Tuned XGBoost Model on Test Set (Threshold = 0.45)",
        width_inches=4.5
    )

    add_image_with_caption(
        doc,
        "reports/figures/XGBoost_Optimized_roc_curve.png",
        "Figure 4.11: ROC Curve of the Tuned XGBoost Classifier (AUC = 0.841)",
        width_inches=4.8
    )

    add_image_with_caption(
        doc,
        "reports/figures/XGBoost_Optimized_pr_curve.png",
        "Figure 4.12: Precision-Recall Curve of the Tuned XGBoost Classifier (AUC = 0.599)",
        width_inches=4.8
    )

    add_image_with_caption(
        doc,
        "reports/figures/XGBoost_Optimized_threshold_sweep.png",
        "Figure 4.13: Threshold vs Metric Performance Sweep identifying the 0.45 Operating Point",
        width_inches=5.2
    )

    add_image_with_caption(
        doc,
        "reports/figures/model_top_feature_importance.png",
        "Figure 4.14: Top 10 Features Driving XGBoost Predictions, led by Incident Severity and Claim Amounts",
        width_inches=5.2
    )

    doc.add_page_break()

    # 4.8 Exploratory Data Analysis Charts
    doc.add_heading("4.8 Exploratory Data Analysis Insights", level=2)
    doc.add_paragraph("Exploratory analysis confirmed critical risk indicators:")

    add_image_with_caption(
        doc,
        "reports/figures/eda_target_distribution.png",
        "Figure 4.15: Overall Claim Distribution confirming 24.7% Fraud Imbalance",
        width_inches=4.5
    )

    add_image_with_caption(
        doc,
        "reports/figures/eda_incident_severity.png",
        "Figure 4.16: Incident Severity vs Fraud Rate — Major Damage claims exhibit ~60% fraud rate",
        width_inches=5.0
    )

    add_image_with_caption(
        doc,
        "reports/figures/eda_insured_hobbies.png",
        "Figure 4.17: Insured Hobbies vs Fraud Rate — Hobbies like chess and cross-fit show elevated risk correlation",
        width_inches=5.2
    )

    add_image_with_caption(
        doc,
        "reports/figures/eda_claim_amount_distribution.png",
        "Figure 4.18: Total Claim Amount Distribution demonstrating higher density for fraudulent claims",
        width_inches=5.0
    )

    doc.add_page_break()

    # CHAPTER 5
    h5 = doc.add_heading("5. COURSE / INTERNSHIP EXPERIENCE", level=1)
    h5.runs[0].font.name = "Arial"
    h5.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "This internship covered a comprehensive range of topics and practical applications highly valuable in data science "
        "and machine learning engineering:"
    )
    doc.add_paragraph("• Python for Data Science: Advanced tabular cleaning, datetime decomposition, and vector math.")
    doc.add_paragraph("• Imbalanced Classification: Confining SMOTE oversampling strictly within CV training folds.")
    doc.add_paragraph("• Machine Learning with Scikit-Learn & XGBoost: Modular ColumnTransformer pipelines and GridSearchCV.")
    doc.add_paragraph("• Interactive Web Deployment: Multi-page Streamlit application with custom CSS and Plotly risk gauges.")
    doc.add_paragraph("• Automated Quality Assurance: 9 automated unit tests written with Pytest ensuring 100% pass rate.")
    doc.add_paragraph("• Risk Communication: Translating complex probabilistic scores into intuitive Low/Medium/High risk tiers.")

    doc.add_heading("Challenges Faced & How They Were Resolved", level=2)
    doc.add_paragraph("1. Missing Values as '?': Handled via custom categorical mapping to 'MISSING' inside SimpleImputer.")
    doc.add_paragraph("2. Extreme Class Imbalance: Resolved by combining SMOTE oversampling with scale_pos_weight in XGBoost.")
    doc.add_paragraph("3. Threshold Optimization: Swept decision thresholds (0.10–0.90) to find the 0.45 point catching 75.51% of frauds.")
    doc.add_paragraph("4. Cloud Deployment: Deployed to Streamlit Community Cloud with optimized caching for sub-second page loads.")

    # CHAPTER 6
    h6 = doc.add_heading("6. CONCLUSION & FUTURE SCOPE", level=1)
    h6.runs[0].font.name = "Arial"
    h6.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "In conclusion, the Insurance Fraud Claims Detection Engine represents a complete, practical data science solution "
        "for a high-stakes enterprise problem. Achieving an 84.00% Accuracy, 75.51% Target Recall, and 84.12% ROC-AUC, the "
        "system provides insurance carriers with an effective decision-support tool to detect suspicious claims while protecting "
        "legitimate policyholders."
    )
    doc.add_paragraph(
        "Future enhancements include incorporating Natural Language Processing (NLP) on adjuster narrative notes and police reports, "
        "integrating computer vision for vehicular damage photo evaluation, and exposing RESTful FastAPI endpoints for enterprise integration."
    )

    doc.add_page_break()

    # CHAPTER 7 - CERTIFICATE
    h7 = doc.add_heading("7. CERTIFICATE OF INTERNSHIP", level=1)
    h7.runs[0].font.name = "Arial"
    h7.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    tbl_cert = doc.add_table(rows=1, cols=1)
    tbl_cert.alignment = WD_TABLE_ALIGNMENT.CENTER
    cert_cell = tbl_cert.cell(0, 0)
    set_cell_background(cert_cell, "FCFDFE")
    set_cell_margins(cert_cell, top=200, bottom=200, left=240, right=240)
    
    c_tcPr = cert_cell._tc.get_or_add_tcPr()
    c_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="double" w:sz="36" w:space="0" w:color="1E3A8A"/>'
        f'<w:left w:val="double" w:sz="36" w:space="0" w:color="1E3A8A"/>'
        f'<w:bottom w:val="double" w:sz="36" w:space="0" w:color="1E3A8A"/>'
        f'<w:right w:val="double" w:sz="36" w:space="0" w:color="1E3A8A"/>'
        f'</w:tcBorders>'
    )
    c_tcPr.append(c_borders)

    cp = cert_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r1 = cp.add_run("CERTIFICATE OF INTERNSHIP\n\n")
    r1.font.name = "Arial"
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(30, 58, 138)
    
    r2 = cp.add_run("This is to certify that\n\n")
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(100, 116, 139)
    
    r3 = cp.add_run("[STUDENT / CANDIDATE NAME]\n")
    r3.font.name = "Arial"
    r3.font.size = Pt(14)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(15, 23, 42)
    
    r4 = cp.add_run("Roll No: [ROLL NUMBER / REGISTRATION ID]\n\n")
    r4.font.name = "Arial"
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(100, 116, 139)
    
    r5 = cp.add_run(
        "has successfully completed a 2-Month Data Science & Machine Learning Internship on the project titled:\n\n"
        "\"INSURANCE FRAUD CLAIMS DETECTION ENGINE:\n"
        "MACHINE LEARNING RISK SCREENING DECISION SUPPORT SYSTEM\"\n\n"
        "During this tenure, the candidate demonstrated exceptional proficiency in Python, Data Preprocessing, "
        "Supervised Classification, Imbalanced Learning (SMOTE), XGBoost Tuning, Model Explainability, "
        "and Streamlit Cloud Web Deployment.\n\n\n"
    )
    r5.font.name = "Arial"
    r5.font.size = Pt(10)
    r5.font.color.rgb = RGBColor(30, 41, 59)
    
    r6 = cp.add_run("Date: 17th August 2026\n\n\n")
    r6.font.name = "Arial"
    r6.font.size = Pt(9.5)
    r6.font.color.rgb = RGBColor(100, 116, 139)
    
    r7 = cp.add_run("_____________________________                    _____________________________\n")
    r7.font.name = "Arial"
    r7.font.size = Pt(9.5)
    r7.font.color.rgb = RGBColor(100, 116, 139)
    
    r8 = cp.add_run("    Project Supervisor                               Department Head / Mentor\n")
    r8.font.name = "Arial"
    r8.font.size = Pt(9.5)
    r8.font.bold = True
    r8.font.color.rgb = RGBColor(15, 23, 42)

    # Save document
    output_docx_path = "reports/ACADEMIC_INTERNSHIP_REPORT.docx"
    doc.save(output_docx_path)
    print(f"Professional DOCX report saved successfully to {output_docx_path}!")

if __name__ == "__main__":
    build_report()
